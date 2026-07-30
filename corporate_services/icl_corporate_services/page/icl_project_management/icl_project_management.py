import calendar

import frappe
from frappe.utils import flt, getdate

from corporate_services.api.project import (
    _get_next_milestones,
    get_last_report_date,
    get_report_frequency_days,
)
from corporate_services.api.project.lifecycle_toolkit import (
    DEFAULT_INTRO_DESCRIPTION,
    DEFAULT_INTRO_TITLE,
    get_project_toolkit_document_template_targets,
    get_project_toolkit_folder_blueprint,
)
from corporate_services.api.project.permissions import get_bypass_roles

PROJECT_PHASE_ORDER = [
    "Initiation",
    "Planning",
    "Execution",
    "Monitoring & Control",
    "Post-Implementation Closeout",
    "Closed",
]


@frappe.whitelist()
def get_dashboard_data():
    summary = _get_summary()
    status_breakdown = _get_status_breakdown()
    projects = _get_projects()
    return {
        "summary": summary,
        "status_breakdown": status_breakdown,
        "projects": projects,
    }


def _split_lines(value):
    if not value:
        return []
    return [line.strip() for line in str(value).splitlines() if line.strip()]


@frappe.whitelist()
def get_lifecycle_config():
    phases = get_project_toolkit_folder_blueprint()

    templates_by_phase = {}
    for template in get_project_toolkit_document_template_targets():
        for placement in template.get("placements") or []:
            phase_name = placement.get("project_phase") or "General"
            names = templates_by_phase.setdefault(phase_name, [])
            if template["document_name"] not in names:
                names.append(template["document_name"])

    for phase in phases:
        phase["templates"] = templates_by_phase.get(phase.get("phase_name"), [])

    return {
        "intro_title": _get_intro_title(),
        "intro_description": _get_intro_description(),
        "phases": phases,
    }


@frappe.whitelist()
def get_template_library():
    rows = frappe.get_all(
        "Project Toolkit Document Templates",
        fields=[
            "name",
            "document_name",
            "description",
            "is_active",
            "attachment",
            "attach_doctype",
            "target_doctype",
        ],
        order_by="document_name asc",
    )
    for row in rows:
        row["requirement"] = row.pop("document_name")
        row["template_file"] = row.pop("attachment")
        target_doctype = row.pop("target_doctype")
        row["doctype"] = target_doctype if row.pop("attach_doctype") else None
    return rows


@frappe.whitelist()
def link_template_file(requirement=None, file_url=None, docname=None):
    docname = (docname or requirement or "").strip()
    if not docname:
        frappe.throw("Requirement is required.")
    if not file_url:
        frappe.throw("Template file is required.")

    if not frappe.db.exists("Project Toolkit Document Templates", docname):
        frappe.throw(f"Template '{docname}' was not found.")

    frappe.db.set_value("Project Toolkit Document Templates", docname, "attachment", file_url)
    return {"name": docname, "template_file": file_url}


def _get_summary():
    rows = frappe.db.sql(
        """
        select
            count(name) as total_projects,
            sum(case when status = 'Completed' then 1 else 0 end) as completed_projects,
            sum(case when status in ('Open', 'Working') then 1 else 0 end) as active_projects,
            avg(coalesce(percent_complete, 0)) as average_progress
        from `tabProject`
        where ifnull(status, '') != 'Cancelled'
        """,
        as_dict=True,
    )
    return rows[0] if rows else {}


def _get_status_breakdown():
    return frappe.db.sql(
        """
        select
            coalesce(status, 'Not Set') as status,
            count(*) as count
        from `tabProject`
        where ifnull(status, '') != 'Cancelled'
        group by coalesce(status, 'Not Set')
        order by count desc
        """,
        as_dict=True,
    )


def _get_projects():
    return frappe.db.sql(
        """
        select
            name,
            project_name,
            status,
            percent_complete,
            expected_start_date,
            expected_end_date,
            priority,
            modified
        from `tabProject`
        where ifnull(status, '') != 'Cancelled'
        order by modified desc
        """,
        as_dict=True,
    )


def _format_month_year(month_year):
    try:
        month, year = month_year.split("-")
        return f"{calendar.month_name[int(month)]} {year}"
    except Exception:
        return month_year


@frappe.whitelist()
def get_project_hours_summary(month_year=None):
    month_year = (month_year or "").strip()

    conditions = ["ts.docstatus != 2", "tpp.project is not null", "tpp.project != ''"]
    params = {}
    if month_year:
        conditions.append("ts.month_year = %(month_year)s")
        params["month_year"] = month_year
    where = " and ".join(conditions)

    project_rows = frappe.db.sql(
        f"""
        select
            tpp.project as project,
            coalesce(p.project_name, tpp.project_name, tpp.project) as project_title,
            sum(tpp.total_hours) as total_hours,
            count(distinct ts.employee) as employee_count
        from `tabTimesheet Submission List` tpp
        inner join `tabTimesheet Submission` ts on ts.name = tpp.parent
        left join `tabProject` p on p.name = tpp.project
        where {where}
        group by tpp.project
        order by total_hours desc
        """,
        params,
        as_dict=True,
    )

    total_hours = sum(flt(row.total_hours) for row in project_rows)
    for row in project_rows:
        row["total_hours"] = flt(row.total_hours)
        row["percentage"] = round((row["total_hours"] / total_hours) * 100, 1) if total_hours else 0

    employee_count = frappe.db.sql(
        f"""
        select count(distinct ts.employee) as c
        from `tabTimesheet Submission List` tpp
        inner join `tabTimesheet Submission` ts on ts.name = tpp.parent
        where {where}
        """,
        params,
        as_dict=True,
    )[0]["c"]

    month_rows = frappe.db.sql(
        """
        select distinct month_year
        from `tabTimesheet Submission`
        where docstatus != 2 and ifnull(month_year, '') != ''
        """,
        as_dict=True,
    )
    months = sorted(
        (m.month_year for m in month_rows),
        key=lambda my: tuple(reversed([int(p) for p in my.split("-")])),
        reverse=True,
    )

    return {
        "total_hours": total_hours,
        "employee_count": employee_count,
        "projects": project_rows,
        "months": [{"value": my, "label": _format_month_year(my)} for my in months],
    }


def _user_is_smt(user):
    if user == "Administrator":
        return True
    return bool(set(frappe.get_roles(user)) & get_bypass_roles())


def _user_employee(user):
    return frappe.db.get_value("Employee", {"user_id": user}, "name")


PROJECT_PHASE_SHORT = {
    "Initiation": "Initiation",
    "Planning": "Planning",
    "Execution": "Execution",
    "Monitoring & Control": "Monitoring",
    "Post-Implementation Closeout": "Closeout",
    "Closed": "Closed",
}

# RAG health status names (not to be confused with the display colors they map to on the frontend).
_RAG_RANK = {"Red": 3, "Amber": 2, "NotStarted": 1, "Green": 0}
_RAG_SUMMARY_KEY = {"Red": "red", "Amber": "amber", "Green": "green", "NotStarted": "not_started"}


def _format_phase(phase, short=False):
    if not phase:
        return None
    if phase not in PROJECT_PHASE_ORDER:
        return phase
    idx = PROJECT_PHASE_ORDER.index(phase) + 1
    if short:
        return f"P{idx} {PROJECT_PHASE_SHORT.get(phase, phase)}"
    return f"Phase {idx} - {phase}"


def _get_open_high_risk_counts(project_names):
    if not project_names:
        return {}
    assessments = frappe.get_all(
        "Project Risk Assessment", filters={"project": ["in", project_names]}, fields=["name", "project"]
    )
    if not assessments:
        return {}
    project_by_assessment = {a.name: a.project for a in assessments}
    risk_rows = frappe.get_all(
        "Risk Assessment List",
        filters={
            "parent": ["in", list(project_by_assessment.keys())],
            "risk_impact": "High",
            "status": ["not in", ["Mitigated", "Closed"]],
        },
        fields=["parent"],
    )
    counts = {}
    for row in risk_rows:
        project = project_by_assessment.get(row.parent)
        if project:
            counts[project] = counts.get(project, 0) + 1
    return counts


def _get_hours_logged(project_names):
    if not project_names:
        return {}
    rows = frappe.db.sql(
        """
        select tpp.project as project, sum(tpp.total_hours) as hours
        from `tabTimesheet Submission List` tpp
        inner join `tabTimesheet Submission` ts on ts.name = tpp.parent
        where ts.docstatus != 2 and tpp.project in %(projects)s
        group by tpp.project
        """,
        {"projects": project_names},
        as_dict=True,
    )
    return {row.project: flt(row.hours) for row in rows}


def _get_task_counts(project_names):
    if not project_names:
        return {}
    rows = frappe.db.sql(
        """
        select
            project,
            count(case when status not in ('Completed','Cancelled','Closed') then 1 end) as open_tasks,
            count(case when status not in ('Completed','Cancelled','Closed')
                        and (status = 'Overdue' or (exp_end_date is not null and exp_end_date < curdate()))
                   then 1 end) as overdue_tasks
        from `tabTask`
        where project in %(projects)s
        group by project
        """,
        {"projects": project_names},
        as_dict=True,
    )
    return {row.project: {"open": row.open_tasks, "overdue": row.overdue_tasks} for row in rows}


def _score_projects(rows):
    """Mutate each project row in place with rag/badge/phase/milestone/hours fields.
    Returns the Red/Amber/Green/NotStarted summary counts.
    """
    project_names = [r.name for r in rows]
    milestones = _get_next_milestones(project_names)
    risk_counts = _get_open_high_risk_counts(project_names)
    hours_by_project = _get_hours_logged(project_names)

    today = getdate()
    summary = {"red": 0, "amber": 0, "green": 0, "not_started": 0}

    for p in rows:
        pct = flt(p.percent_complete)
        end_date = getdate(p.expected_end_date) if p.expected_end_date else None

        freq_days = get_report_frequency_days(p.get("report_frequency"))
        report_badge = None
        if freq_days:
            last_report = get_last_report_date(p.name)
            if not last_report:
                report_badge = {"type": "no_report", "text": "No status report yet"}
            else:
                overdue_days = (today - getdate(last_report)).days - freq_days
                if overdue_days > 0:
                    report_badge = {
                        "type": "report_overdue",
                        "text": f"Report {overdue_days} day{'s' if overdue_days != 1 else ''} overdue",
                    }

        risk_count = risk_counts.get(p.name, 0)

        if pct <= 0:
            rag = "NotStarted"
        elif risk_count > 0 or (end_date and end_date < today):
            rag = "Red"
        elif report_badge or (end_date and 0 <= (end_date - today).days <= 14):
            rag = "Amber"
        else:
            rag = "Green"

        badge = None
        if risk_count > 0:
            badge = {"type": "risk", "text": f"{risk_count} High risk{'s' if risk_count != 1 else ''}"}
        elif report_badge:
            badge = report_badge

        milestone = milestones.get(p.name)

        summary[_RAG_SUMMARY_KEY[rag]] += 1

        p["percent_complete"] = pct
        p["rag"] = rag
        p["phase_raw"] = p.get("phase")
        p["phase"] = _format_phase(p.get("phase"))
        p["badge"] = badge
        p["next_milestone"] = (milestone or {}).get("subject")
        p["next_milestone_date"] = (milestone or {}).get("exp_end_date")
        p["days_remaining"] = (end_date - today).days if end_date else None
        p["hours_logged"] = hours_by_project.get(p.name, 0)
        p.pop("report_frequency", None)

    return summary


def _get_active_pms():
    # Not filtering by project status for now - show every PM with a project assignment.
    return frappe.db.sql(
        """
        select distinct pm.employee, pm.employee_name
        from `tabProject Manager` pm
        inner join `tabProject` p on p.name = pm.parent
        where pm.employee is not null
        order by pm.employee_name asc
        """,
        as_dict=True,
    )


def _get_pm_breakdown(employee=None):
    # Not filtering by project status for now - show every project regardless of status.
    if employee:
        rows = frappe.db.sql(
            """
            select
                p.name,
                p.project_name,
                coalesce(p.percent_complete, 0) as percent_complete,
                p.expected_end_date,
                p.custom_project_phase as phase,
                p.custom_project_report_frequency as report_frequency
            from `tabProject` p
            inner join `tabProject Manager` pm on pm.parent = p.name and pm.employee = %(employee)s
            order by p.expected_end_date asc
            """,
            {"employee": employee},
            as_dict=True,
        )
    else:
        rows = frappe.db.sql(
            """
            select distinct
                p.name,
                p.project_name,
                coalesce(p.percent_complete, 0) as percent_complete,
                p.expected_end_date,
                p.custom_project_phase as phase,
                p.custom_project_report_frequency as report_frequency
            from `tabProject` p
            order by p.expected_end_date asc
            """,
            as_dict=True,
        )
    _score_projects(rows)

    project_names = [r.name for r in rows]
    task_counts = _get_task_counts(project_names)

    projects = []
    overdue_reports = 0
    open_tasks_total = 0
    for r in rows:
        counts = task_counts.get(r.name, {"open": 0, "overdue": 0})
        open_tasks_total += counts["open"]
        if r["badge"] and r["badge"]["type"] in ("no_report", "report_overdue"):
            overdue_reports += 1

        milestone_label = None
        if r["next_milestone"]:
            due_label = frappe.utils.formatdate(r["next_milestone_date"], "MMM d") if r["next_milestone_date"] else None
            milestone_label = f"{r['next_milestone']}{' ' + due_label if due_label else ''}"

        projects.append({
            "project": r.name,
            "project_name": r.project_name,
            "phase": _format_phase(r.get("phase_raw"), short=True),
            "open_tasks": counts["open"],
            "overdue_tasks": counts["overdue"],
            "next_milestone": milestone_label,
            "rag": r["rag"],
        })

    return {
        "projects": projects,
        "active_projects": len(projects),
        "open_tasks": open_tasks_total,
        "overdue_reports": overdue_reports,
    }


def _get_all_pm_comparison(pm_rows):
    rank_to_rag = {v: k for k, v in _RAG_RANK.items()}
    comparison = []
    for pm_row in pm_rows:
        breakdown = _get_pm_breakdown(pm_row.employee)
        overdue_tasks = sum(p["overdue_tasks"] for p in breakdown["projects"])
        worst_rank = max((_RAG_RANK[p["rag"]] for p in breakdown["projects"]), default=0)
        comparison.append({
            "employee": pm_row.employee,
            "employee_name": pm_row.employee_name,
            "active_projects": breakdown["active_projects"],
            "open_tasks": breakdown["open_tasks"],
            "overdue_tasks": overdue_tasks,
            "overdue_reports": breakdown["overdue_reports"],
            "health": rank_to_rag[worst_rank],
        })

    comparison.sort(key=lambda c: (-_RAG_RANK[c["health"]], -c["overdue_tasks"]))
    return comparison


@frappe.whitelist()
def get_portfolio_health(pm=None):
    user = frappe.session.user
    is_smt = _user_is_smt(user)

    pm_rows = _get_active_pms()

    conditions = ["p.status not in ('Cancelled', 'Completed')"]
    params = {}

    if is_smt:
        if pm:
            conditions.append(
                "exists (select 1 from `tabProject Manager` pmf where pmf.parent = p.name and pmf.employee = %(pm)s)"
            )
            params["pm"] = pm
    else:
        employee = _user_employee(user)
        if not employee:
            return {
                "projects": [],
                "summary": {"red": 0, "amber": 0, "green": 0, "not_started": 0},
                "is_smt": False,
                "pms": [],
                "selected_pm": None,
            }
        conditions.append(
            "exists (select 1 from `tabProject Manager` pmf where pmf.parent = p.name and pmf.employee = %(employee)s)"
        )
        params["employee"] = employee

    where = " and ".join(conditions)

    rows = frappe.db.sql(
        f"""
        select
            p.name,
            p.project_name,
            p.status,
            coalesce(p.percent_complete, 0) as percent_complete,
            p.expected_start_date,
            p.expected_end_date,
            p.customer,
            p.custom_project_phase as phase,
            p.custom_project_report_frequency as report_frequency,
            group_concat(distinct pm.employee_name separator ', ') as pm_names
        from `tabProject` p
        left join `tabProject Manager` pm on pm.parent = p.name
        where {where}
        group by p.name
        order by p.expected_end_date asc
        """,
        params,
        as_dict=True,
    )

    summary = _score_projects(rows)
    for p in rows:
        p.pop("phase_raw", None)

    return {
        "projects": rows,
        "summary": summary,
        "is_smt": is_smt,
        "pms": pm_rows,
        "selected_pm": pm if is_smt else None,
    }


@frappe.whitelist()
def get_delivery_pipeline():
    stages = frappe.db.sql(
        """
        select distinct stage_name, display_order
        from `tabHIS Project Lifecycle Stage`
        where is_active = 1
        order by display_order asc, stage_name asc
        """,
        as_dict=True,
    )
    seen, unique_stages = set(), []
    for s in stages:
        if s.stage_name not in seen:
            seen.add(s.stage_name)
            unique_stages.append(s.stage_name)
    if not unique_stages:
        unique_stages = ["Prepare", "Plan", "Design", "Development", "Implementation", "Maintenance"]

    projects = frappe.db.sql(
        """
        select
            p.name,
            p.project_name,
            p.status,
            coalesce(p.percent_complete, 0) as percent_complete,
            p.expected_start_date,
            p.expected_end_date,
            p.customer,
            group_concat(pm.employee_name separator ', ') as pm_names
        from `tabProject` p
        left join `tabProject Manager` pm on pm.parent = p.name
        where p.status not in ('Cancelled')
        group by p.name
        order by p.expected_end_date asc
        """,
        as_dict=True,
    )
    n = len(unique_stages)
    for proj in projects:
        pct = float(proj.get("percent_complete") or 0)
        completed = int((pct / 100) * n)
        proj["stage_progress"] = [
            "complete" if i < completed else ("current" if i == completed else "pending")
            for i in range(n)
        ]
    return {"projects": projects, "stages": unique_stages}


@frappe.whitelist()
def get_overdue_deliverables():
    rows = frappe.db.sql(
        """
        select
            t.name,
            t.subject,
            t.project,
            t.exp_end_date,
            t.status,
            datediff(curdate(), t.exp_end_date) as days_overdue,
            p.project_name,
            p.customer,
            group_concat(pm.employee_name separator ', ') as pm_names
        from `tabTask` t
        join `tabProject` p on p.name = t.project
        left join `tabProject Manager` pm on pm.parent = t.project
        where t.project is not null
          and t.status not in ('Completed', 'Cancelled', 'Closed')
          and (
              t.status = 'Overdue'
              or (t.exp_end_date is not null and t.exp_end_date < curdate())
          )
          and p.status not in ('Cancelled', 'Completed')
        group by t.name
        order by days_overdue desc
        limit 100
        """,
        as_dict=True,
    )
    return {"tasks": rows, "total": len(rows)}


@frappe.whitelist()
def get_pm_workload(pm=None):
    user = frappe.session.user
    is_smt = _user_is_smt(user)
    pm_rows = _get_active_pms()

    if is_smt:
        # SMT sees every PM's projects by default; `pm` narrows to one PM's breakdown.
        target_employee = pm or None
    else:
        target_employee = _user_employee(user)
        if not target_employee:
            return {
                "is_smt": False,
                "employee": None,
                "employee_name": None,
                "pms": [],
                "my_view": None,
                "smt_view": None,
            }

    if target_employee:
        employee_name = frappe.db.get_value("Employee", target_employee, "employee_name")
    elif is_smt:
        employee_name = "All PMs"
    else:
        employee_name = None

    my_view = _get_pm_breakdown(target_employee)

    smt_view = _get_all_pm_comparison(pm_rows) if is_smt else None

    return {
        "is_smt": is_smt,
        "employee": target_employee,
        "employee_name": employee_name,
        "pms": pm_rows,
        "my_view": my_view,
        "smt_view": smt_view,
    }


@frappe.whitelist()
def get_lessons_learned_trends():
    workflow_states = frappe.db.sql(
        """
        select
            coalesce(nullif(workflow_state, ''), 'Draft') as state,
            count(*) as count
        from `tabProject Management Lessons Learned`
        group by coalesce(nullif(workflow_state, ''), 'Draft')
        order by count desc
        """,
        as_dict=True,
    )

    recommendation_priorities = frappe.db.sql(
        """
        select
            coalesce(nullif(trim(r.priority), ''), 'Not set') as priority,
            count(*) as count
        from `tabLessons Learned Recommendation` r
        join `tabProject Management Lessons Learned` ll on ll.name = r.parent
        where ll.workflow_state = 'Approved'
        group by coalesce(nullif(trim(r.priority), ''), 'Not set')
        order by field(r.priority, 'High', 'Medium', 'Low') asc, count desc
        """,
        as_dict=True,
    )

    next_step_statuses = frappe.db.sql(
        """
        select
            coalesce(nullif(trim(ns.status), ''), 'Not set') as status,
            count(*) as count
        from `tabLessons Learned Next Step` ns
        join `tabProject Management Lessons Learned` ll on ll.name = ns.parent
        where ll.workflow_state = 'Approved'
        group by coalesce(nullif(trim(ns.status), ''), 'Not set')
        order by count desc
        """,
        as_dict=True,
    )

    coverage = frappe.db.sql(
        """
        select
            sum(case when exists (
                select 1 from `tabLessons Learned Root Cause` rc where rc.parent = ll.name
            ) then 1 else 0 end) as with_root_causes,
            sum(case when exists (
                select 1 from `tabLessons Learned Recommendation` r where r.parent = ll.name
            ) then 1 else 0 end) as with_recommendations,
            sum(case when exists (
                select 1 from `tabLessons Learned Next Step` ns where ns.parent = ll.name
            ) then 1 else 0 end) as with_next_steps,
            count(*) as total
        from `tabProject Management Lessons Learned` ll
        where ll.workflow_state = 'Approved'
        """,
        as_dict=True,
    )

    root_causes = frappe.db.sql(
        """
        select
            rc.issue,
            rc.root_cause,
            rc.area_affected,
            ll.name as report_name,
            ll.project_title,
            ll.reporter_name
        from `tabLessons Learned Root Cause` rc
        join `tabProject Management Lessons Learned` ll on ll.name = rc.parent
        where ll.workflow_state = 'Approved'
          and (rc.issue is not null or rc.root_cause is not null)
        order by ll.name, rc.idx
        """,
        as_dict=True,
    )

    recommendations = frappe.db.sql(
        """
        select
            r.recommendation,
            r.priority,
            r.area,
            ll.name as report_name,
            ll.project_title,
            ll.reporter_name
        from `tabLessons Learned Recommendation` r
        join `tabProject Management Lessons Learned` ll on ll.name = r.parent
        where ll.workflow_state = 'Approved'
          and r.recommendation is not null
        order by field(r.priority, 'High', 'Medium', 'Low'), ll.name, r.idx
        """,
        as_dict=True,
    )

    next_steps = frappe.db.sql(
        """
        select
            ns.action_item,
            ns.responsible_person,
            ns.deadline,
            ns.status,
            ll.name as report_name,
            ll.project_title
        from `tabLessons Learned Next Step` ns
        join `tabProject Management Lessons Learned` ll on ll.name = ns.parent
        where ll.workflow_state = 'Approved'
          and ns.action_item is not null
        order by ns.deadline asc, ll.name, ns.idx
        """,
        as_dict=True,
    )

    return {
        "workflow_states": workflow_states,
        "recommendation_priorities": recommendation_priorities,
        "next_step_statuses": next_step_statuses,
        "coverage": coverage[0] if coverage else {},
        "root_causes": root_causes,
        "recommendations": recommendations,
        "next_steps": next_steps,
    }


@frappe.whitelist()
def search_lessons_learned_kb(q=None, area=None, priority=None, project_type=None):
    q = (q or "").strip()
    area = (area or "").strip()
    priority = (priority or "").strip()
    project_type = (project_type or "").strip()

    conditions = ["ll.workflow_state = 'Approved'"]
    params = {}

    if project_type:
        conditions.append(
            """exists (
                select 1 from `tabProject` p
                where p.name = ll.project and p.project_type = %(project_type)s
            )"""
        )
        params["project_type"] = project_type

    if q:
        conditions.append(
            """(
                ll.project_title like %(q)s
                or exists (
                    select 1 from `tabLessons Learned Root Cause` rc
                    where rc.parent = ll.name
                      and (rc.issue like %(q)s or rc.root_cause like %(q)s)
                )
                or exists (
                    select 1 from `tabLessons Learned Recommendation` r
                    where r.parent = ll.name
                      and r.recommendation like %(q)s
                )
                or exists (
                    select 1 from `tabLessons Learned Next Step` ns
                    where ns.parent = ll.name
                      and ns.action_item like %(q)s
                )
            )"""
        )
        params["q"] = f"%{q}%"

    if area:
        conditions.append(
            """exists (
                select 1 from `tabLessons Learned Recommendation` r
                where r.parent = ll.name and r.area = %(area)s
            )"""
        )
        params["area"] = area

    if priority:
        conditions.append(
            """exists (
                select 1 from `tabLessons Learned Recommendation` r
                where r.parent = ll.name and r.priority = %(priority)s
            )"""
        )
        params["priority"] = priority

    where = " and ".join(conditions)
    reports = frappe.db.sql(
        f"""
        select ll.name as report_name, ll.project_title, ll.reporter_name,
               ll.workflow_state, ll.date_of_report
        from `tabProject Management Lessons Learned` ll
        where {where}
        order by ll.date_of_report desc
        limit 50
        """,
        params,
        as_dict=True,
    )

    names = [r["report_name"] for r in reports]
    if not names:
        return []

    fmt = ",".join(["%s"] * len(names))

    root_causes = frappe.db.sql(
        f"select parent, issue, root_cause from `tabLessons Learned Root Cause` where parent in ({fmt}) order by idx",
        names, as_dict=True,
    )
    recommendations = frappe.db.sql(
        f"select parent, recommendation, priority, area from `tabLessons Learned Recommendation` where parent in ({fmt}) order by idx",
        names, as_dict=True,
    )
    next_steps = frappe.db.sql(
        f"select parent, action_item, responsible_person, deadline, status from `tabLessons Learned Next Step` where parent in ({fmt}) order by idx",
        names, as_dict=True,
    )

    rc_map = {}
    for rc in root_causes:
        rc_map.setdefault(rc["parent"], []).append(rc)
    rec_map = {}
    for rec in recommendations:
        rec_map.setdefault(rec["parent"], []).append(rec)
    ns_map = {}
    for ns in next_steps:
        ns_map.setdefault(ns["parent"], []).append(ns)

    for r in reports:
        n = r["report_name"]
        r["root_causes"] = rc_map.get(n, [])
        r["recommendations"] = rec_map.get(n, [])
        r["next_steps"] = ns_map.get(n, [])

    return reports


@frappe.whitelist()
def export_lessons_learned_docx(report_name=None):
    if not report_name:
        frappe.throw("report_name is required.")

    doc = frappe.get_doc("Project Management Lessons Learned", report_name)

    root_causes = frappe.db.get_all(
        "Lessons Learned Root Cause",
        filters={"parent": report_name},
        fields=["issue", "root_cause", "area_affected"],
        order_by="idx",
    )
    recommendations = frappe.db.get_all(
        "Lessons Learned Recommendation",
        filters={"parent": report_name},
        fields=["recommendation", "priority", "area"],
        order_by="idx",
    )
    next_steps = frappe.db.get_all(
        "Lessons Learned Next Step",
        filters={"parent": report_name},
        fields=["action_item", "responsible_person", "deadline", "status"],
        order_by="idx",
    )
    answers = frappe.db.get_all(
        "Lessons Learned Answer",
        filters={"parent": report_name},
        fields=["question", "response"],
        order_by="idx",
    )

    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import io, os

    word = Document()

    # Page margins
    for section in word.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    BRAND = RGBColor(0x00, 0x52, 0x9B)   # IntelliSOFT blue

    def heading(text, level=1):
        p = word.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13 if level == 1 else 11)
        run.font.color.rgb = BRAND
        return p

    def body(text):
        p = word.add_paragraph(str(text or "-"))
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(10)
        return p

    def add_table(headers, rows):
        tbl = word.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        hrow = tbl.rows[0]
        for i, h in enumerate(headers):
            cell = hrow.cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shading = cell._tc.get_or_add_tcPr()
            from docx.oxml import OxmlElement
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "00529B")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:val"), "clear")
            shading.append(shd)
        for row_data in rows:
            row = tbl.add_row()
            for i, val in enumerate(row_data):
                row.cells[i].text = str(val or "-")
                row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        word.add_paragraph()

    # ── Header ──
    p = word.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INTELLISOFT CONSULTING")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND

    p2 = word.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Lessons Learned Report")
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    word.add_paragraph()

    # ── Project Info ──
    heading("Project Information")
    info_tbl = word.add_table(rows=0, cols=2)
    info_tbl.style = "Table Grid"
    for label, value in [
        ("Report No.", doc.name),
        ("Project", doc.project_title or doc.project or "-"),
        ("Client", getattr(doc, "client", None) or "-"),
        ("Reporter", doc.reporter_name or "-"),
        ("Designation", getattr(doc, "designation_of_reporter", None) or "-"),
        ("Date of Report", str(doc.date_of_report or "-")),
        ("Expected Start", str(getattr(doc, "expected_start_date", None) or "-")),
        ("Expected End", str(getattr(doc, "expected_end_date", None) or "-")),
        ("Workflow Status", doc.workflow_state or "-"),
    ]:
        row = info_tbl.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = str(value)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    word.add_paragraph()

    # ── Narrative ──
    if answers:
        heading("Narrative Report")
        for ans in answers:
            p = word.add_paragraph()
            run = p.add_run(f"{ans.get('question', '')}:")
            run.bold = True
            run.font.size = Pt(10)
            from html.parser import HTMLParser
            class _Strip(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.parts = []
                def handle_data(self, d):
                    self.parts.append(d)
            s = _Strip()
            s.feed(str(ans.get("response") or ""))
            body("".join(s.parts).strip() or "-")

    # ── Root Cause Analysis ──
    if root_causes:
        heading("Root Cause Analysis")
        add_table(
            ["Issue", "Root Cause", "Area Affected"],
            [(r.issue, r.root_cause, r.area_affected) for r in root_causes],
        )

    # ── Recommendations ──
    if recommendations:
        heading("Recommendations")
        add_table(
            ["Recommendation", "Priority", "Area"],
            [(r.recommendation, r.priority, r.area) for r in recommendations],
        )

    # ── Next Steps ──
    if next_steps:
        heading("Next Steps / Follow-up")
        add_table(
            ["Action Item", "Responsible Person", "Deadline", "Status"],
            [(n.action_item, n.responsible_person, str(n.deadline or ""), n.status) for n in next_steps],
        )

    # ── Footer note ──
    word.add_paragraph()
    p = word.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IntelliSOFT Consulting - Project Management Office - Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Save ──
    buf = io.BytesIO()
    word.save(buf)
    buf.seek(0)

    fname = f"LL-{report_name}.docx"
    file_doc = frappe.get_doc({
        "doctype": "File",
        "file_name": fname,
        "is_private": 1,
        "content": buf.read(),
    })
    file_doc.save(ignore_permissions=True)

    return {"file_url": file_doc.file_url}


def _get_intro_title():
    doc = _get_lifecycle_doc()
    if doc and getattr(doc, "intro_title", None):
        return doc.intro_title
    return DEFAULT_INTRO_TITLE


def _get_intro_description():
    doc = _get_lifecycle_doc()
    if doc and getattr(doc, "intro_description", None):
        return doc.intro_description
    return DEFAULT_INTRO_DESCRIPTION


def _get_lifecycle_doc():
    if not frappe.db.exists("DocType", "HIS Project Lifecycle Config"):
        return None
    try:
        return frappe.get_single("HIS Project Lifecycle Config")
    except Exception:
        return None
